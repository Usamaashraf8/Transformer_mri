"""
Run this script once to generate the project documentation Word file.
    python generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
    return heading


def add_code_block(doc, code):
    para = doc.add_paragraph()
    run  = para.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    para.paragraph_format.left_indent = Inches(0.4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F0F0F0")
    para._element.get_or_add_pPr().append(shading)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
    doc.add_paragraph()


def build_doc():
    doc = Document()

    # ── Title page ────────────────────────────────────────────────────────────
    title = doc.add_heading("Brain Tumor MRI Classification", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Vision Transformer (ViT) — Built From Scratch")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)

    doc.add_paragraph("Dataset: Training 5600 images | Testing 1600 images | 4 Classes")
    doc.add_page_break()

    # ── Table of Contents (manual) ────────────────────────────────────────────
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Project Overview",
        "2. Dataset Description",
        "3. Project Structure",
        "4. File-by-File Documentation",
        "   4.1  config.py",
        "   4.2  dataset.py",
        "   4.3  model.py",
        "   4.4  train.py",
        "   4.5  evaluate.py",
        "   4.6  main.py",
        "5. Model Architecture Deep Dive",
        "6. Training Strategy",
        "7. How to Run",
        "8. Expected Outputs",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Bullet" if item.startswith(" ") else "Normal")
    doc.add_page_break()

    # ── 1. Project Overview ───────────────────────────────────────────────────
    add_heading(doc, "1. Project Overview")
    doc.add_paragraph(
        "This project implements a Vision Transformer (ViT) entirely from scratch using PyTorch "
        "to classify brain MRI scans into four categories: Glioma, Meningioma, Pituitary Tumor, "
        "and No Tumor. No pre-trained weights are used — every parameter is randomly initialised "
        "and learned solely from the provided dataset. "
        "The project follows a clean modular structure so that each concern (data, model, "
        "training, evaluation) is kept in its own file."
    )

    # ── 2. Dataset Description ────────────────────────────────────────────────
    add_heading(doc, "2. Dataset Description")
    doc.add_paragraph(
        "The dataset is organized in a standard ImageFolder layout under the 'dataset/' directory."
    )
    add_table(
        doc,
        headers=["Split", "Glioma", "Meningioma", "Pituitary", "No Tumor", "Total"],
        rows=[
            ["Training", "1400", "1400", "1400", "1400", "5600"],
            ["Testing",  "400",  "400",  "400",  "400",  "1600"],
            ["Total",    "1800", "1800", "1800", "1800", "7200"],
        ],
    )
    doc.add_paragraph(
        "The dataset is perfectly balanced (equal samples per class), which means no class "
        "weighting is needed in the loss function. Images are JPEG files of varying sizes "
        "and are resized to 224×224 pixels during preprocessing."
    )
    doc.add_paragraph("Class descriptions:")
    class_desc = [
        ("Glioma",     "A tumor that starts in the glial cells of the brain or spine."),
        ("Meningioma", "A tumor that forms on the membranes surrounding the brain and spinal cord."),
        ("Pituitary",  "A tumor that forms in the pituitary gland at the base of the brain."),
        ("No Tumor",   "Healthy brain MRI scan with no detectable tumor."),
    ]
    for name, desc in class_desc:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    # ── 3. Project Structure ──────────────────────────────────────────────────
    add_heading(doc, "3. Project Structure")
    add_code_block(doc,
        "Transformer_mri/\n"
        "├── dataset/\n"
        "│   ├── Training/\n"
        "│   │   ├── glioma/        (1400 images)\n"
        "│   │   ├── meningioma/    (1400 images)\n"
        "│   │   ├── notumor/       (1400 images)\n"
        "│   │   └── pituitary/     (1400 images)\n"
        "│   └── Testing/\n"
        "│       ├── glioma/        (400 images)\n"
        "│       ├── meningioma/    (400 images)\n"
        "│       ├── notumor/       (400 images)\n"
        "│       └── pituitary/     (400 images)\n"
        "├── checkpoints/           (created at runtime)\n"
        "│   ├── best_model.pth\n"
        "│   ├── training_curves.png\n"
        "│   └── confusion_matrix.png\n"
        "├── config.py\n"
        "├── dataset.py\n"
        "├── model.py\n"
        "├── train.py\n"
        "├── evaluate.py\n"
        "├── main.py\n"
        "├── generate_doc.py\n"
        "└── .gitignore"
    )

    # ── 4. File-by-File ───────────────────────────────────────────────────────
    add_heading(doc, "4. File-by-File Documentation")

    # 4.1 config.py
    add_heading(doc, "4.1  config.py", level=2)
    doc.add_paragraph(
        "Purpose: Central configuration file that stores every hyperparameter and path used "
        "across the entire project. Having a single config file ensures that changing a value "
        "(e.g., batch size or learning rate) in one place automatically propagates to all modules. "
        "No magic numbers are scattered through the codebase."
    )
    doc.add_paragraph("Key settings defined:")
    add_table(
        doc,
        headers=["Parameter", "Value", "Meaning"],
        rows=[
            ["IMAGE_SIZE",     "224",    "All images are resized to 224×224"],
            ["PATCH_SIZE",     "16",     "Each patch is 16×16 pixels → 196 patches per image"],
            ["EMBED_DIM",      "256",    "Dimension of each token embedding vector"],
            ["NUM_HEADS",      "8",      "Number of parallel attention heads"],
            ["DEPTH",          "6",      "Number of stacked Transformer Encoder blocks"],
            ["MLP_DIM",        "1024",   "Hidden size of the feed-forward network (4×EMBED_DIM)"],
            ["DROPOUT",        "0.1",    "Dropout probability applied throughout the model"],
            ["BATCH_SIZE",     "32",     "Number of images per training step"],
            ["EPOCHS",         "50",     "Total number of training passes over the dataset"],
            ["LEARNING_RATE",  "1e-3",   "Initial learning rate for AdamW"],
            ["WEIGHT_DECAY",   "0.05",   "L2 regularisation coefficient in AdamW"],
            ["VAL_SPLIT",      "0.1",    "10% of training data reserved for validation"],
        ],
    )

    # 4.2 dataset.py
    add_heading(doc, "4.2  dataset.py", level=2)
    doc.add_paragraph(
        "Purpose: Handles all data-related operations — loading images from disk, applying "
        "preprocessing transforms, splitting training data into train/val sets, and exposing "
        "PyTorch DataLoader objects for use during training and evaluation."
    )
    doc.add_paragraph("Why it is required:")
    doc.add_paragraph(
        "PyTorch requires a Dataset subclass that implements __len__ and __getitem__. "
        "Keeping this logic separate from training keeps the code clean and reusable. "
        "The DataLoader adds batching, shuffling, and multi-process data loading for speed.",
        style="List Bullet",
    )
    doc.add_paragraph("Key components:")
    components = [
        ("BrainTumorDataset",       "Custom Dataset class. Walks class folders, builds a list of (path, label) pairs."),
        ("get_train_transforms()",  "Augmentation pipeline: resize, random flip, random rotation, color jitter, normalize to [-1, 1]."),
        ("get_val_test_transforms()","No augmentation — only resize and normalize (ensures fair evaluation)."),
        ("get_dataloaders()",        "Splits training set 90/10, returns three DataLoaders: train, val, test."),
    ]
    for name, desc in components:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    # 4.3 model.py
    add_heading(doc, "4.3  model.py", level=2)
    doc.add_paragraph(
        "Purpose: Defines the full Vision Transformer architecture from scratch using "
        "only PyTorch primitives (nn.Linear, nn.LayerNorm, etc.). No external model "
        "libraries or pre-trained weights are used."
    )
    doc.add_paragraph("Why ViT (Vision Transformer)?")
    doc.add_paragraph(
        "Traditional CNNs (e.g., ResNet) capture local features through convolution. "
        "ViT instead treats an image as a sequence of patches and uses self-attention "
        "to model global relationships between any two patches — regardless of distance. "
        "This makes it well-suited for detecting spatially distributed tumor patterns.",
        style="List Bullet",
    )
    doc.add_paragraph("Building blocks:")
    blocks = [
        ("PatchEmbedding",         "Splits 224×224 image into 196 patches of size 16×16, then projects each to a 256-dim vector using Conv2d."),
        ("MultiHeadSelfAttention", "Computes scaled dot-product attention across all 197 tokens (196 patches + 1 CLS) using 8 parallel heads."),
        ("MLP",                    "Two-layer feed-forward network (256→1024→256) with GELU activation — applied independently to each token."),
        ("TransformerBlock",       "Pre-LayerNorm residual block: LN → Attention → residual; LN → MLP → residual. Repeated 6 times."),
        ("VisionTransformer",      "Full pipeline: patch embed → prepend CLS token → add positional embeddings → 6 blocks → LayerNorm → CLS head → 4 logits."),
    ]
    for name, desc in blocks:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)
    doc.add_paragraph(
        "The [CLS] token is a special learnable vector prepended to the patch sequence. "
        "After all transformer blocks, only the CLS token's output is fed to the "
        "classification head. This is the standard ViT classification mechanism."
    )

    # 4.4 train.py
    add_heading(doc, "4.4  train.py", level=2)
    doc.add_paragraph(
        "Purpose: Implements the full training loop — forward pass, loss computation, "
        "backpropagation, optimisation, learning-rate scheduling, and model checkpointing."
    )
    doc.add_paragraph("Why each component:")
    choices = [
        ("CrossEntropyLoss",       "Standard multi-class classification loss. Combines log-softmax and NLL loss in one numerically stable operation."),
        ("AdamW optimizer",        "Adam with decoupled weight decay. Better generalisation than vanilla Adam, especially for transformers."),
        ("CosineAnnealingLR",      "Smoothly reduces LR from 1e-3 to ~0 over 50 epochs. Avoids the sharp LR drops of step schedules and helps find better minima."),
        ("Val accuracy checkpoint","Saves the model only when validation accuracy improves, preventing saving an overfit model."),
        ("plot_curves()",          "Saves loss and accuracy curves to disk for post-training analysis without needing a notebook."),
    ]
    for name, desc in choices:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    # 4.5 evaluate.py
    add_heading(doc, "4.5  evaluate.py", level=2)
    doc.add_paragraph(
        "Purpose: Loads the best saved model weights and evaluates on the held-out test set. "
        "Produces a detailed per-class report and a confusion matrix image."
    )
    doc.add_paragraph(
        "The test set is never touched during training. Running evaluate.py separately "
        "ensures a clean, unbiased measurement of generalisation performance. "
        "The confusion matrix visually shows which tumor types are most confused "
        "with each other — useful diagnostic information for medical AI systems."
    )

    # 4.6 main.py
    add_heading(doc, "4.6  main.py", level=2)
    doc.add_paragraph(
        "Purpose: Single entry point that runs the entire pipeline — training then evaluation — "
        "with a single command: python main.py. It imports and calls train.main() and evaluate.main() "
        "in sequence. Useful for end-to-end runs on a remote machine or scheduled job."
    )

    # ── 5. Architecture Deep Dive ─────────────────────────────────────────────
    add_heading(doc, "5. Model Architecture Deep Dive")
    doc.add_paragraph("The forward pass through the ViT proceeds as follows:")
    steps = [
        "Input image: (Batch, 3, 224, 224)",
        "PatchEmbedding (Conv2d 16×16): (B, 196, 256)",
        "Prepend CLS token: (B, 197, 256)",
        "Add positional embedding (197×256 learnable): (B, 197, 256)",
        "Transformer Block ×6 — each block:",
        "    a. LayerNorm → Multi-Head Self-Attention (8 heads, head_dim=32) → residual add",
        "    b. LayerNorm → MLP (256→1024→256, GELU) → residual add",
        "Final LayerNorm: (B, 197, 256)",
        "Extract CLS token: (B, 256)",
        "Linear classification head: (B, 4)",
        "Output: logits for [glioma, meningioma, notumor, pituitary]",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number" if not s.startswith(" ") else "List Bullet")

    doc.add_paragraph(f"Total trainable parameters: ~4.2 million")

    # ── 6. Training Strategy ──────────────────────────────────────────────────
    add_heading(doc, "6. Training Strategy")
    doc.add_paragraph("The following choices are made specifically for training from scratch on a small dataset:")
    strategy = [
        ("Data Augmentation",    "Random flips, rotations, and color jitter artificially expand the effective dataset size and improve generalisation."),
        ("Weight Initialisation","Truncated normal (std=0.02) for all linear layers. This is the standard ViT initialisation."),
        ("Pre-LN architecture",  "LayerNorm before (not after) attention/MLP. More stable gradients when training from scratch."),
        ("Dropout (0.1)",        "Applied in attention, MLP, and positional embedding. Prevents co-adaptation of neurons."),
        ("AdamW + weight decay", "Decoupled L2 regularisation prevents overfitting on 5600 training samples."),
        ("Cosine LR Annealing",  "Gradually reduces learning rate, allowing fine-grained convergence in later epochs."),
        ("Small model",          "EMBED_DIM=256, DEPTH=6 keeps the model expressiveness-vs-data ratio reasonable."),
    ]
    for name, desc in strategy:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    # ── 7. How to Run ─────────────────────────────────────────────────────────
    add_heading(doc, "7. How to Run")
    doc.add_paragraph("Step 1: Install dependencies")
    add_code_block(doc, "pip install torch torchvision matplotlib scikit-learn python-docx")
    doc.add_paragraph("Step 2: Run full pipeline (train + evaluate)")
    add_code_block(doc, "python main.py")
    doc.add_paragraph("Step 3: Or run separately")
    add_code_block(doc, "python train.py       # training only\npython evaluate.py    # evaluation only (requires best_model.pth)")
    doc.add_paragraph("Note: Change DEVICE in config.py to 'cuda' if you have an NVIDIA GPU, "
                      "or 'mps' for Apple Silicon Mac (default).")

    # ── 8. Expected Outputs ───────────────────────────────────────────────────
    add_heading(doc, "8. Expected Outputs")
    add_table(
        doc,
        headers=["File", "Location", "Description"],
        rows=[
            ["best_model.pth",       "checkpoints/", "Model weights with highest validation accuracy"],
            ["training_curves.png",  "checkpoints/", "Loss and accuracy plots over all epochs"],
            ["confusion_matrix.png", "checkpoints/", "Per-class prediction matrix on test set"],
        ],
    )
    doc.add_paragraph(
        "Console output during training shows epoch-by-epoch train/val loss and accuracy. "
        "After evaluation, a full sklearn classification_report is printed with per-class "
        "precision, recall, and F1 score."
    )

    path = "/Users/ashraf/Desktop/Transformer_mri/ViT_BrainTumor_Documentation.docx"
    doc.save(path)
    print(f"Documentation saved to: {path}")


if __name__ == "__main__":
    build_doc()
